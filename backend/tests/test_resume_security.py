import pytest
import io
import docx

def get_token(client, username):
    response = client.post(
        "/auth/token",
        data={"username": username, "password": "password123"}
    )
    return response.json()["access_token"]

def create_mock_docx():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@example.com")
    doc.add_paragraph("Skills: Python, Go, Docker")
    doc.add_paragraph("Experience: 5 years")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def test_resume_type_restriction(client):
    token = get_token(client, "recruiter_user")
    
    file_data = b"fake-jpg-content"
    response = client.post(
        "/upload_resume",
        files={"file": ("resume.jpg", file_data, "image/jpeg")},
        headers={"Authorization": f"Bearer {token}"}
    )
    assert response.status_code == 400
    assert "Unsupported file format" in response.json()["detail"]

def test_resume_size_restriction(client):
    token = get_token(client, "recruiter_user")
    
    large_data = b"0" * (5 * 1024 * 1024 + 100)
    response = client.post(
        "/upload_resume",
        files={"file": ("resume.pdf", large_data, "application/pdf")},
        headers={"Authorization": f"Bearer {token}"}
    )
    assert response.status_code == 400
    assert "exceeds" in response.json()["detail"]

def test_filename_sanitization_traversal(client):
    token = get_token(client, "recruiter_user")
    
    file_data = b"Jane Doe resume data"
    response = client.post(
        "/upload_resume",
        files={"file": ("../../../../etc/passwd.txt", file_data, "text/plain")},
        headers={"Authorization": f"Bearer {token}"}
    )
    assert response.status_code == 201
    assert "name" in response.json()

def test_docx_parsing(client):
    token = get_token(client, "recruiter_user")
    
    docx_bytes = create_mock_docx()
    response = client.post(
        "/upload_resume",
        files={"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Authorization": f"Bearer {token}"}
    )
    assert response.status_code == 201
    assert "name" in response.json()
